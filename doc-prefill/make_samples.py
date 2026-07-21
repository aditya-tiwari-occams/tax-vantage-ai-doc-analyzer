"""Generate realistic synthetic R&D project documents (PDF + DOCX) for extraction testing.
Content mirrors what an Occams TaxVantage client would upload to describe a 2025 R&D project.
"""
import os
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, ListFlowable, ListItem)
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
SAMP = os.path.join(HERE, "samples")
os.makedirs(SAMP, exist_ok=True)

# ---- Shared content: two projects in one document (multi-project case) ----
PROJECTS = [
    {
        "title": "Adaptive Cold-Chain Routing Engine",
        "contract": "Work was not done under Contract",
        "description": (
            "In 2025 our engineering team developed a new routing engine to keep "
            "perishable pharmaceutical shipments within a 2-8 degC window across "
            "multi-leg journeys. The existing third-party routers optimized only for "
            "time and cost and could not incorporate real-time reefer telemetry, so we "
            "built a custom system from the ground up."
        ),
        "uncertainty": (
            "At project start it was technically uncertain whether we could predict "
            "in-transit temperature excursions accurately enough to re-route shipments "
            "before spoilage, given noisy IoT sensor data and intermittent connectivity "
            "in transit. We did not know which modeling approach would generalize across "
            "carriers, packaging types, and climate zones."
        ),
        "experimentation": (
            "We ran a systematic process of experimentation: (1) prototyped three "
            "forecasting approaches - a physics-based thermal model, a gradient-boosted "
            "tree, and an LSTM; (2) evaluated each against 14 months of historical "
            "excursion data; (3) iterated on feature engineering for sensor gap-filling; "
            "and (4) A/B tested the winning model against the legacy router on live lanes."
        ),
        "alternatives": (
            "Alternatives considered: licensing an off-the-shelf telematics SaaS "
            "(rejected - no re-routing API); extending our existing rules engine "
            "(rejected - could not model nonlinear thermal dynamics); and a pure "
            "physics simulation (rejected - too slow for real-time decisions)."
        ),
        "man_hours": 4200,
        "employees": 6,
        "supplies": "AWS GPU compute, IoT reefer sensors, test packaging units",
    },
    {
        "title": "Self-Healing Warehouse Vision QA",
        "contract": "Fixed Price",
        "description": (
            "We designed a computer-vision quality-assurance station that detects "
            "damaged packaging on a high-speed conveyor. The novel aspect was achieving "
            "reliable detection at 180 units/minute under variable warehouse lighting "
            "without halting the line for recalibration."
        ),
        "uncertainty": (
            "It was uncertain whether a single model could maintain >99% precision as "
            "lighting, product SKUs, and conveyor speed varied, and whether on-device "
            "inference latency could stay under 40ms on the available edge hardware."
        ),
        "experimentation": (
            "Process of experimentation included evaluating four model architectures, "
            "synthetic data augmentation for rare defect classes, quantization "
            "experiments to hit the latency budget, and a shadow-mode deployment "
            "comparing model calls against human inspectors over six weeks."
        ),
        "alternatives": (
            "We evaluated manual inspection (too slow, inconsistent), a laser "
            "profilometry rig (prohibitively expensive and SKU-specific), and a "
            "cloud-inference design (rejected due to latency and bandwidth)."
        ),
        "man_hours": 3100,
        "employees": 4,
        "supplies": "Edge GPU devices, industrial cameras, annotation tooling",
    },
]

def build_pdf(path):
    doc = SimpleDocTemplate(path, pagesize=LETTER,
                            topMargin=0.8*inch, bottomMargin=0.8*inch)
    ss = getSampleStyleSheet()
    h1 = ss["Heading1"]; h2 = ss["Heading2"]; body = ss["BodyText"]
    body.spaceAfter = 8
    story = []
    story.append(Paragraph("Acme Logistics, Inc. — 2025 R&D Project Documentation", h1))
    story.append(Paragraph("Prepared for R&D Tax Credit Study (Tax Year 2025)", body))
    story.append(Spacer(1, 10))
    for i, p in enumerate(PROJECTS, 1):
        story.append(Paragraph(f"Project {i}: {p['title']}", h2))
        # a small facts table to test table extraction + reading order
        tbl = Table([
            ["Contract Type", p["contract"]],
            ["Total Man Hours", str(p["man_hours"])],
            ["Employees on Research", str(p["employees"])],
            ["Supplies Used", p["supplies"]],
        ], colWidths=[2.0*inch, 4.0*inch])
        tbl.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 8))
        story.append(Paragraph("<b>Project Description</b>", body))
        story.append(Paragraph(p["description"], body))
        story.append(Paragraph("<b>Technical Uncertainty</b>", body))
        story.append(Paragraph(p["uncertainty"], body))
        story.append(Paragraph("<b>Process of Experimentation</b>", body))
        story.append(Paragraph(p["experimentation"], body))
        story.append(Paragraph("<b>Alternatives Considered</b>", body))
        story.append(Paragraph(p["alternatives"], body))
        story.append(Spacer(1, 14))
    doc.build(story)
    print("PDF written:", path)

def build_docx(path):
    d = Document()
    d.add_heading("Beacon Software LLC — R&D Narrative (2025)", level=1)
    p = PROJECTS[0]  # single-project Word doc
    d.add_heading(f"Project: {p['title']}", level=2)
    # order-sensitive: table appears BEFORE some paragraphs
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in [("Contract Type", p["contract"]),
                 ("Total Man Hours", str(p["man_hours"])),
                 ("Employees on Research", str(p["employees"])),
                 ("Supplies Used", p["supplies"])]:
        row = t.add_row().cells
        row[0].text = k; row[1].text = v
    d.add_heading("Project Description", level=3)
    d.add_paragraph(p["description"])
    d.add_heading("Technical Uncertainty", level=3)
    d.add_paragraph(p["uncertainty"])
    d.add_heading("Process of Experimentation", level=3)
    d.add_paragraph(p["experimentation"])
    d.add_heading("Alternatives Considered", level=3)
    d.add_paragraph(p["alternatives"])
    d.save(path)
    print("DOCX written:", path)

if __name__ == "__main__":
    build_pdf(os.path.join(SAMP, "acme_two_projects.pdf"))
    build_docx(os.path.join(SAMP, "beacon_one_project.docx"))
